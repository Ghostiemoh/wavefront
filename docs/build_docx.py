"""Build Wavefront.docx from ARTICLE.md, embedding referenced images."""
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches, RGBColor

HERE = Path(__file__).parent
ARTICLE = HERE / "ARTICLE.md"
OUT = HERE / "Wavefront.docx"

raw = ARTICLE.read_text(encoding="utf-8")

doc = Document()

# Page setup — narrow margins for readability
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

# Base styles
styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

for style_name, size, bold in [
    ("Heading 1", 24, True),
    ("Heading 2", 18, True),
    ("Heading 3", 14, True),
]:
    s = styles[style_name]
    s.font.name = "Calibri"
    s.font.size = Pt(size)
    s.font.bold = bold
    s.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[[^\]]+\]\([^)]+\))"
)


def add_inline(paragraph, text):
    """Add a markdown-formatted run sequence to an existing paragraph."""
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            r = paragraph.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = paragraph.add_run(part[1:-1])
            r.italic = True
        elif part.startswith("`") and part.endswith("`"):
            r = paragraph.add_run(part[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        elif part.startswith("[") and "](" in part:
            label = part[1 : part.index("]")]
            url = part[part.index("(") + 1 : -1]
            add_hyperlink(paragraph, url, label)
        else:
            paragraph.add_run(part)


def add_hyperlink(paragraph, url, text):
    part = paragraph.part
    rId = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rId)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2563EB")
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_image(path, max_width_in=6.5):
    """Embed an image, sized to fit the page width."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(max_width_in))


def parse_table(lines, start):
    """Parse a markdown table starting at index `start`. Returns (rows, end_idx)."""
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        if set(line.replace("|", "").replace("-", "").replace(":", "").strip()) == set():
            i += 1
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append(cells)
        i += 1
    return rows, i


lines = raw.splitlines()
i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    if not stripped:
        i += 1
        continue

    # Image
    img_match = re.match(r"!\[(.*?)\]\((.+?)\)", stripped)
    if img_match:
        alt, path = img_match.group(1), img_match.group(2)
        img_path = HERE / path
        if img_path.exists():
            add_image(img_path)
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cap.add_run(alt)
            r.italic = True
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)
        else:
            p = doc.add_paragraph()
            r = p.add_run(f"[missing image: {path}]")
            r.italic = True
            r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        i += 1
        continue

    # Horizontal rule
    if stripped == "---":
        doc.add_paragraph().add_run().add_break()
        i += 1
        continue

    # Headings
    if stripped.startswith("# "):
        h = doc.add_heading(stripped[2:], level=1)
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        i += 1
        continue
    if stripped.startswith("## "):
        doc.add_heading(stripped[3:], level=2)
        i += 1
        continue
    if stripped.startswith("### "):
        doc.add_heading(stripped[4:], level=3)
        i += 1
        continue

    # Blockquote
    if stripped.startswith("> "):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run(stripped[2:])
        r.italic = True
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        # also process inline markdown by adjusting — keep simple
        i += 1
        continue

    # Table
    if stripped.startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
        rows, new_i = parse_table(lines, i)
        if rows:
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Light Grid Accent 1"
            for ri, row in enumerate(rows):
                for ci, cell_text in enumerate(row):
                    if ci >= len(table.rows[ri].cells):
                        continue
                    cell = table.rows[ri].cells[ci]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    add_inline(p, cell_text)
                    if ri == 0:
                        for run in p.runs:
                            run.bold = True
        i = new_i
        continue

    # Unordered list
    if re.match(r"^[-*]\s+", stripped):
        while i < len(lines) and re.match(r"^[-*]\s+", lines[i].strip()):
            item = re.sub(r"^[-*]\s+", "", lines[i].strip())
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, item)
            i += 1
        continue

    # Ordered list
    if re.match(r"^\d+\.\s+", stripped):
        while i < len(lines) and re.match(r"^\d+\.\s+", lines[i].strip()):
            item = re.sub(r"^\d+\.\s+", "", lines[i].strip())
            p = doc.add_paragraph(style="List Number")
            add_inline(p, item)
            i += 1
        continue

    # Plain paragraph (collect consecutive non-empty non-special lines)
    buf = []
    while i < len(lines):
        s = lines[i].strip()
        if not s:
            break
        if s.startswith(("#", ">", "|", "!", "- ", "* ")) or re.match(r"^\d+\.\s+", s) or s == "---":
            break
        buf.append(s)
        i += 1
    if buf:
        para = doc.add_paragraph()
        add_inline(para, " ".join(buf))

doc.save(OUT)
print(f"Wrote {OUT} ({OUT.stat().st_size:,} bytes)")
